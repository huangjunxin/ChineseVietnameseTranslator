import gradio as gr
from docx import Document
import fitz  # PyMuPDF
from datetime import datetime
import os
from dotenv import load_dotenv

from app.routes.text_translator import translate_text

load_dotenv()
passcode_key = os.getenv("PASSCODE_KEY")


def process_uploaded_file(file):
    content = ""
    # Detect the file type based on its extension
    file_extension = file.name.split('.')[-1].lower()

    # Process each file according to its type
    if file_extension == 'docx':
        doc = Document(file)
        full_text = [para.text for para in doc.paragraphs]
        content = '\n'.join(full_text)
    elif file_extension == 'pdf':
        # Open the PDF file
        with fitz.open(file) as pdf:
            full_text = []
            for page in pdf:
                full_text.append(page.get_text("text"))
            content = '\n'.join(full_text)
    else:
        raise ValueError(f"Unsupported file type: {file_extension}")

    return content


def translate_document(source_language, target_language, original_doc, tone_of_voice, industry, model, passcode):
    output_files = []

    # Check if the passcode is correct
    if passcode + "1" != passcode_key:
        # Create a new .docx document
        doc = Document()
        doc.add_paragraph("The passcode is incorrect. Please try again.")

        # Save the document
        timestamp = datetime.now().strftime("%m%d%H%M")
        output_filename = f"translated_text_{timestamp}.docx"
        doc.save(output_filename)
        output_files.append(output_filename)

        # Return the list of output files
        return output_files

    # Generate translated text
    # Process the uploaded files
    original_doc_content = ""

    if original_doc is not None:
        original_doc_content = process_uploaded_file(original_doc)

    # Translate the uploaded files
    translated_text = translate_text(source_language, target_language, original_doc_content, tone_of_voice, industry,
                                     model, passcode)

    # Create a new .docx document
    doc = Document()
    doc.add_paragraph(f"{translated_text}")

    # Save the document
    timestamp = datetime.now().strftime("%m%d%H%M")
    output_filename = f"translated_text_{timestamp}.docx"
    doc.save(output_filename)
    output_files.append(output_filename)

    # Return the list of output files
    return output_files


# Interface for Document Translator
document_translator = gr.Interface(
    fn=translate_document,
    inputs=[
        gr.Dropdown(
            label="Source Language",
            choices=["Chinese", "English (UK)", "English (US)", "Vietnamese", "Japanese", "Korean", "French", "German",
                     "Spanish", "Portuguese (Brazilian)", "Portuguese (European)", "Italian", "Dutch", "Polish",
                     "Russian"],
            value="Chinese"
        ),
        gr.Dropdown(
            label="Target Language",
            choices=["Chinese", "English (UK)", "English (US)", "Vietnamese", "Japanese", "Korean", "French", "German",
                     "Spanish", "Portuguese (Brazilian)", "Portuguese (European)", "Italian", "Dutch", "Polish",
                     "Russian"],
            value="Vietnamese"
        ),
        gr.File(
            label="Original Text Document",
            file_types=[".docx", ".pdf"],
            file_count="single"
        ),
        gr.Radio(
            label="Tone of Voice",
            choices=["Standard", "Formal", "Informal"],
            value="Standard"
        ),
        gr.Dropdown(
            label="Industry Sector",
            choices=["General Fields", "Academic Papers", "Biomedicine", "Information Technology",
                     "Finance and Economics", "News and Information", "Aerospace", "Mechanical Manufacturing",
                     "Laws and Regulations", "Humanities and Social Sciences"],
            value="General Fields"
        ),
        gr.Dropdown(
            label="Model Provider (Model Name)",
            choices=["DeepL", "Volcengine", "HKBU ChatGPT (gpt-35-turbo)", "HKBU ChatGPT (gpt-4-turbo)",
                     "OpenAI (gpt-3.5-turbo-0125)", "OpenAI (gpt-4-turbo-preview)", "Google Gemini (gemini-pro)",
                     "Baichuan AI (Baichuan2)", "Zhipu AI (glm-3-turbo)", "Zhipu AI (glm-4)"],
            value="OpenAI (gpt-3.5-turbo-1106)"
        ),
        gr.Textbox(
            label="Passcode",
            placeholder="Enter the passcode here",
            type="password",
            lines=1,
            max_lines=1
        )
    ],
    outputs=[
        gr.File(label="Translated Text Document", file_count="single", type="filepath")
    ],
    title="FloweryTranslator - Document Translator"
)
